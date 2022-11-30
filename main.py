import xlsxwriter
import docx
import os


workbook = xlsxwriter.Workbook('CopyOfVizits.xlsx')
worksheet = workbook.add_worksheet()
worksheet.set_column('A:A', 30)
worksheet.write('A2', 'Копия Визитки 1 сторона:')
worksheet.insert_image('B2', 'vizitka.png')
worksheet.write('A12','Копия Визитки 2 сторона:')
worksheet.insert_image('B12', 'Vizitka2.png', {'x_offset': 15, 'y_offset': 10})
workbook.close()
os.startfile('CopyOfVizits.xlsx')

doc = docx.Document()

doc.add_paragraph('Копия Визитки 1 сторона')
doc.add_picture('vizitka.png', width = docx.shared.Cm(10))
doc.add_paragraph('Копия Визитки 2 сторона')
doc.add_picture('vizitka2.png', width = docx.shared.Cm(10))
doc.save('CopyOfVizits.docx')
os.startfile('CopyOfVizits.docx')
